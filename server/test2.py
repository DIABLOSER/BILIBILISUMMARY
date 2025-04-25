#根据视频文件识别语音转文字并生成总结文档
import whisper
from moviepy.editor import VideoFileClip
import os
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from docx import Document
import noisereduce as nr
import soundfile as sf
import numpy as np
def extract_audio_with_moviepy(video_path, audio_path="temp_audio.wav"):
    # 用 moviepy 提取音频
    clip = VideoFileClip(video_path)
    clip.audio.write_audiofile(audio_path, codec='pcm_s16le', fps=16000, nbytes=2, buffersize=2000, ffmpeg_params=["-ac", "1"])
    return audio_path

def reduce_noise(audio_path, output_path="temp_audio_denoised.wav"):
    # 加载音频文件
    data, rate = sf.read(audio_path)
    # 降噪处理
    reduced_noise = nr.reduce_noise(y=data, sr=rate)
    # 保存降噪后的音频
    sf.write(output_path, reduced_noise, rate)
    return output_path

def transcribe_audio(audio_path, model_name="medium"):
    # 使用更高级的模型进行语音识别
    model = whisper.load_model(model_name)
    result = model.transcribe(audio_path, language="zh")
    return result["text"]

def generate_summary(text):
    # 使用 sumy 生成总结
    parser = PlaintextParser.from_string(text, Tokenizer("chinese"))
    summarizer = LsaSummarizer()
    summary = summarizer(parser.document, 3)  # 生成3句总结
    return [str(sentence) for sentence in summary]

def save_summary_to_docx(text, summary, output_path="summary.docx"):
    # 创建 Word 文档
    doc = Document()
    doc.add_heading('视频识别结果', 0)
    doc.add_paragraph('识别文本：')
    doc.add_paragraph(text)
    doc.add_paragraph('总结：')
    for sentence in summary:
        doc.add_paragraph(sentence)
    doc.save(output_path)

def video_to_text(video_path):
    audio_path = extract_audio_with_moviepy(video_path)
    denoised_audio_path = reduce_noise(audio_path)  # 降噪处理
    text = transcribe_audio(denoised_audio_path)  # 使用更高级的模型进行识别
    summary = generate_summary(text)
    summary_path = f"summaries/{os.path.basename(video_path).split('.')[0]}.docx"
    save_summary_to_docx(text, summary, summary_path)
    os.remove(audio_path)  # 清理临时文件
    os.remove(denoised_audio_path)  # 清理降噪后的临时文件
    return text, summary, summary_path

if __name__ == "__main__":
    video_file = "server/temp/114256381089136.mp4"  # 替换成你的视频路径
    text, summary, summary_path = video_to_text(video_file)
    print("识别结果：", text)
    print("总结：", summary)
    print("文档路径：", summary_path)